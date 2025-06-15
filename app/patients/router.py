from fastapi import APIRouter, UploadFile, Depends, HTTPException, Body
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm import selectinload

from starlette import status
from starlette.responses import FileResponse, Response, StreamingResponse

from docx import Document
import io

from app.database import SessionDep, async_session_maker
from app.exception import IncorrectEmailOrPasswordException, UserNotFindException
from app.patients.auth import get_password_hash, authenticate_user, create_access_token
from app.patients.dao import PatientDAO
from app.patients.dependencies import get_current_user
from app.patients.models import Patient
from app.patients.schemas import SPatientAdd, SPatient, SPatientAuth, SPatientHosp
from app.patients.utils import (
    recognize_qr_code,
    generate_qr_code,
    generate_consent,
    generate_contract,
)

router = APIRouter(
    prefix="/patients",
    tags=["patients"],
)


@router.post("/")
async def add_patient(session: SessionDep, data_patient: SPatientAdd) -> SPatient:
    user = await PatientDAO.find_one_or_none(session=session, email=data_patient.email)
    if user:
        raise HTTPException(
            status_code=status.HTTP_409_CONFLICT, detail="Пользователь уже существует"
        )

    hashed_password = get_password_hash(data_patient.password)
    data_patient.password = hashed_password

    async with async_session_maker() as session:
        patient = await PatientDAO.add(session, **data_patient.model_dump())
    return patient


@router.get("/all")
async def all_patients(session: SessionDep) -> list[SPatient]:
    patients = await PatientDAO.find_all(session)
    return patients


@router.get("/profile/{patient_id}")
async def get_patient(patient_id: int, session: SessionDep) -> SPatient:
    patient = await PatientDAO.find_one_or_none_by_id(
        session=session, model_id=patient_id
    )
    if not patient:
        raise UserNotFindException
    return patient


@router.post("/login")
async def login_user(
    response: Response, patient_data: SPatientAuth, session: SessionDep
):
    patient = await authenticate_user(
        patient_data.email, patient_data.password, session
    )
    if not patient:
        raise IncorrectEmailOrPasswordException
    access_token = create_access_token({"sub": str(patient.id)})
    response.set_cookie("access_token", access_token, httponly=True)
    return access_token


@router.get("/me")
async def get_me(patient=Depends(get_current_user)) -> SPatient:
    return patient


@router.get("/download/{patient_id}")
async def get_patient_med_card(session: SessionDep, patient_id: int):
    patient = await PatientDAO.find_one_or_none_by_id(session, patient_id)
    if not patient:
        raise HTTPException(status_code=404, detail="Patient not found")

    # Создаем документ Word в памяти
    buffer = io.BytesIO()
    doc = Document()

    # Добавляем заголовок
    doc.add_heading("Медицинская карта пациента", 0)

    # Добавляем данные пациента
    doc.add_paragraph(
        f"ФИО: {patient.last_name} {patient.first_name} {patient.middle_name or ''}"
    )
    doc.add_paragraph(f"Дата рождения: {patient.date_birthday}")
    doc.add_paragraph(f"Пол: {patient.gender.value}")
    doc.add_paragraph(f"Адрес: {patient.address}")
    doc.add_paragraph(f"Телефон: {patient.phone_number}")
    doc.add_paragraph(f"Email: {patient.email}")
    doc.add_paragraph(f"Паспорт: {patient.passport}")
    doc.add_paragraph(f"Номер страхового полиса: {patient.number_insurance_policy}")
    doc.add_paragraph(f"Страховая компания: {patient.insurance_company}")
    doc.add_paragraph(f"Дата выдачи: {patient.date_issue}")
    doc.add_paragraph(f"Дата истечения: {patient.date_expiration}")
    doc.add_paragraph(f"Диагноз: {patient.diagnosis or 'Не установлен'}")
    doc.add_paragraph(f"Дата последнего обращения: {patient.date_last_request}")
    doc.add_paragraph(f"Дата следующего визита: {patient.date_next_visit}")

    # Сохраняем документ в буфер
    doc.save(buffer)
    buffer.seek(0)

    # Возвращаем файл как поток
    return StreamingResponse(
        iter([buffer.getvalue()]),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f"attachment; filename=med_card_{patient_id}.docx"
        },
    )


@router.get("/hospitalizations/{patient_id}")
async def get_all_hosp_by_patient_id(
    session: SessionDep, patient_id: int
) -> SPatientHosp:
    hosp = await PatientDAO.find_one_or_none_by_id(
        session, patient_id, options=[selectinload(Patient.hospitalizations)]
    )
    return hosp


@router.post("/upload_photo/{patient_id}")
async def upload_photo(
    uploaded_file: UploadFile,
    patient_id: int,
    session: SessionDep,
    patient=Depends(get_patient),
):
    """Добавить сохранение в s3 хранилище"""
    print("start")
    file = uploaded_file.file
    file_name = f"app/static/photo/ph_{str(patient_id)}.jpg"
    with open(file_name, "wb") as f:
        f.write(file.read())

    patient.photo_url = file_name

    await session.commit()
    return file_name


@router.get("/qr_code/{patient_id}")
async def get_qr_code_patient(
    patient_id: int, session: SessionDep, patient=Depends(get_patient)
):
    qr_url = await generate_qr_code(str(patient_id))

    patient.qr_code_url = qr_url
    await session.commit()
    return FileResponse(path=qr_url, filename="qr_code.jpg", media_type="image/jpeg")


@router.post("/qr_code")
async def recognition_qr_code_patient(uploaded_file: UploadFile):
    file_content = await uploaded_file.read()
    data = await recognize_qr_code(file_content)
    return {"data": data}


@router.get("/consent/{patient_id}")
async def consent(patient_id: int, session: SessionDep):
    patient = await PatientDAO.find_one_or_none_by_id(session, model_id=patient_id)
    file_path = await generate_consent(patient)

    return FileResponse(
        path=file_path,
        filename="Согласие_на_обработку_данных.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


@router.get("/contract/{patient_id}")
async def contract(patient_id: int, session: SessionDep):
    patient = await PatientDAO.find_one_or_none_by_id(session, model_id=patient_id)

    file_path = await generate_contract(patient)

    return FileResponse(
        path=file_path,
        filename="Договор.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
